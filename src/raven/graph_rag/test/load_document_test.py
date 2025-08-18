import logging
import os
import subprocess
import tempfile
from typing import Any

from docx import Document


def load_document(file_obj: Any) -> str:
    """
    Extract text content from DOCX and DOC files.
    DOC files are converted to DOCX via LibreOffice before extraction.

    Args:
        file_obj: File path (str) or file-like object

    Returns:
        Extracted text content (str)
    """
    # Configure logging (only once)
    if not logging.getLogger().hasHandlers():
        logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s", filename="document_extractor.log")

    try:
        # Determine file type and validate
        filename = getattr(file_obj, "name", "") if not isinstance(file_obj, str) else file_obj
        ext = os.path.splitext(filename)[1].lower()

        if isinstance(file_obj, str) and not os.path.exists(file_obj):
            raise FileNotFoundError(f"File not found: {file_obj}")

        # Process DOCX files directly
        if ext == ".docx":
            if isinstance(file_obj, str):
                doc = Document(file_obj)
            else:
                if hasattr(file_obj, "seek"):
                    file_obj.seek(0)  # Reset file pointer
                doc = Document(file_obj)
            return "\n".join([p.text for p in doc.paragraphs])

        # Process DOC files via LibreOffice conversion
        elif ext == ".doc":
            with tempfile.TemporaryDirectory() as temp_dir:
                # Save input file to temporary location
                input_path = os.path.join(temp_dir, "input.doc")
                if isinstance(file_obj, str):
                    with open(file_obj, "rb") as f_src, open(input_path, "wb") as f_dst:
                        f_dst.write(f_src.read())
                else:
                    file_obj.seek(0)
                    with open(input_path, "wb") as f:
                        f.write(file_obj.read())

                # Build full path to soffice.exe
                libreoffice_exe = os.path.join("C:\\", "Program Files", "LibreOffice", "program", "soffice.exe")

                # Verify executable exists
                if not os.path.exists(libreoffice_exe):
                    raise FileNotFoundError(f"LibreOffice executable not found: {libreoffice_exe}")

                try:
                    # Execute conversion command
                    cmd = [libreoffice_exe, "--headless", "--convert-to", "docx", "--outdir", temp_dir, input_path]

                    subprocess.run(cmd, capture_output=True, text=True, check=True)

                    # Check for converted file
                    converted_files = [f for f in os.listdir(temp_dir) if f.endswith(".docx")]
                    if not converted_files:
                        raise RuntimeError("LibreOffice conversion produced no output")

                    # Load converted DOCX
                    docx_path = os.path.join(temp_dir, converted_files[0])
                    doc = Document(docx_path)
                    return "\n".join([p.text for p in doc.paragraphs])

                except subprocess.CalledProcessError as e:
                    logging.error(f"LibreOffice execution failed: {e.stderr}")
                    raise RuntimeError(f"Conversion failed: {e.stderr}") from e

        else:
            raise ValueError(f"Unsupported file type: {ext}")

    except Exception as e:
        logging.error(f"Document processing failed: {str(e)}")
        raise
