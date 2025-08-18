import json
import logging
import os
import re
import subprocess
import tempfile
from tempfile import NamedTemporaryFile
from typing import Optional, Type, TypeVar, Union, Any
import urllib.parse

from bs4 import BeautifulSoup
from docx import Document
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

T = TypeVar("T")  # 泛型类型变量，代表任意类


class UnsupportedFormatError(Exception):
    pass


class DocumentLoader:
    def __init__(self, listen: Optional[object] = None):
        """
        Initialize DocumentLoader with optional listener

        :param listen: Optional listener object (e.g. MultiAgentListen)
        """
        self.listen = listen

    def _load_json_file(self, source: Union[str, bytes, os.PathLike], target_class: Type[T]) -> T:
        """
        Load JSON file into object
        :param source: file path
        :param target_class: object class
        :return: object
        """
        try:
            # Validate file existence
            if not os.path.exists(source):
                raise FileNotFoundError(f"File not found: {source}")  # type: ignore

            # Read and parse JSON file
            with open(source, "r", encoding="utf-8") as file:
                data = json.load(file)

            # Instantiate target class from JSON data
            if isinstance(data, dict):
                # Create object from dictionary
                return target_class(**data)
            elif isinstance(data, list) and issubclass(target_class, list):
                # Create list subclass instance
                return target_class(data)  # type: ignore
            else:
                raise ValueError(f"Cannot convert JSON data to {target_class.__name__} type")

        except json.JSONDecodeError as e:
            raise ValueError(f"JSON parsing error: {e}") from e
        except Exception as e:
            raise RuntimeError(f"Error loading JSON file: {e}") from e
        pass

    def load(self, source: Union[str, bytes, os.PathLike]) -> str:
        """
        Unified loader for both local files and URLs

        :param source: File path or URL string
        :return: Extracted document text
        """
        if isinstance(source, (bytes, os.PathLike)):
            source = os.fspath(source)

        if source.startswith(("http://", "https://")):  # type: ignore
            return self._load_from_url(source)  # type: ignore
        else:
            return self._load_from_file(source)  # type: ignore

    def _load_from_file(self, file_path: Union[str, os.PathLike]) -> str:
        """
        Load document from local file path

        :param file_path: Path to local file
        :return: Document text
        """
        file_path = os.fspath(file_path)
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = os.path.splitext(file_path)[1][1:]  # 去掉前面的点

        with open(file_path, "rb") as f:
            self._check_file_size(f)
            return self._load_text_from_file(f, file_extension)

    def _get_extension_from_path(self, path: str) -> tuple[str, str]:
        """
        Extract file extension from local path

        :param path: File path
        :return: Lowercase file extension without dot
        """
        return "", ""

    def _load_from_url(self, url: str) -> str:
        """
        Load document from URL

        :param url: HTTP/HTTPS URL pointing to the document
        :return: Document text
        """
        file_extension = self._get_extension_from_url(url)
        temp_file = self._down_load_to_tempfile(url)
        try:
            # Pass the temporary file object directly
            self._check_file_size(temp_file)
            return self._load_text_from_file(temp_file, file_extension)
        finally:
            self._cleanup_temp_file(temp_file)

    def _load_text_from_file(self, file_obj: Any, extension: Optional[str] = None) -> str:
        """
        Route file to appropriate loader based on extension

        :param file_obj: File-like object
        :param extension: File format extension
        :raises UnsupportedFormatError: For unhandled file types
        """
        # Ensure consistent lowercase handling
        extension = extension.lower() if extension else ""

        if extension == "pdf":
            return self._load_pdf(file_obj)
        elif extension in ("docx", "doc"):
            return self._load_docx(file_obj)
        elif extension == "txt":
            return self._load_plain_text(file_obj)
        elif extension == "html":
            return self._load_html(file_obj)
        elif extension == "md":
            return self._load_markdown(file_obj)
        elif extension == "json":
            return self._load_json(file_obj)
        else:
            raise UnsupportedFormatError(f"Unsupported format: {extension}")

    def _check_file_size(self, file_obj: Any) -> bool:
        """
        Validate file size against environment limits

        :param file_obj: File-like object
        """
        # Implementation should:
        # 1. Read environment configuration for size limits
        # 2. Compare file size to limits
        # 3. Raise exception if exceeds limit
        return True

    def _down_load_to_tempfile(self, url: str) -> Any:
        """
        Download remote resource to temporary file

        :param url: Valid document URL
        :return: Temporary file object
        """
        # 1. Make HTTP request with timeout
        session = requests.Session()
        retries = Retry(total=3, backoff_factor=1, status_forcelist=[500, 502, 503, 504])
        session.mount("http://", HTTPAdapter(max_retries=retries))
        session.mount("https://", HTTPAdapter(max_retries=retries))

        # 2. Stream content to temporary file
        temp_file = NamedTemporaryFile(delete=False)
        temp_path = temp_file.name
        temp_file.close()  # Close to allow writing by another process

        try:
            response = session.get(url, stream=True, timeout=30)
            response.raise_for_status()

            with open(temp_path, "wb") as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:  # filter out keep-alive new chunks
                        f.write(chunk)

            # 3. Return open tempfile handle1
            return open(temp_path, "rb")

        except Exception:
            self._cleanup_temp_file(temp_path)
            raise

    def _get_extension_from_url(self, url: str) -> str:
        """
        Extract file extension from URL path.
        :param url: Document URL
        :return: File extension without dot
        """
        # 移除查询参数
        parsed_url = urllib.parse.urlparse(url)
        clean_url = parsed_url.path

        # 处理伪装扩展名（如 malicious.exe.pdf → 取最后一段）
        filename = os.path.basename(clean_url)

        # 支持类型：txt, pdf, doc, docx, html, md, json
        valid_exts = ["txt", "pdf", "doc", "docx", "html", "md", "json"]

        # 获取扩展名（带点号的后缀）
        ext_match = re.search(r"\.([a-zA-Z0-9]+)$", filename)
        if not ext_match:
            return ""
        ext = ext_match.group(1).lower()

        return ext if ext in valid_exts else ""

    def _cleanup_temp_file(self, temp_file: Any) -> None:
        """
        Clean up temporary file resources

        :param temp_file: Temporary file object to clean up
        """
        try:
            file_path = temp_file.name  # type: ignore
            temp_file.close()  # type: ignore
            if os.path.exists(file_path):
                os.unlink(file_path)
        except Exception:
            # Avoid raising exception during cleanup
            pass

    # Implementations of specific file loaders remain unchanged1
    def _load_pdf(self, file_obj: Any) -> str:
        """PDF document loader"""
        return ""

    def _load_docx(self, file_obj: Any) -> str:
        """Word document loader"""
        # Configure logging (only once)
        if not logging.getLogger().hasHandlers():
            logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s", filename="document_extractor.log")

        try:
            # Determine file type and validate
            filename = getattr(file_obj, "name", "") if not isinstance(file_obj, str) else file_obj
            ext = os.path.splitext(filename)[1].lower()

            if isinstance(file_obj, str) and not os.path.exists(file_obj):
                raise FileNotFoundError(f"File not found: {file_obj}")

            # Process DOCX files directly
            if ext == ".docx":
                if isinstance(file_obj, str):
                    doc = Document(file_obj)
                else:
                    if hasattr(file_obj, "seek"):
                        file_obj.seek(0)  # Reset file pointer
                    doc = Document(file_obj)
                return "\n".join([p.text for p in doc.paragraphs])

            # Process DOC files via LibreOffice conversion
            elif ext == ".doc":
                with tempfile.TemporaryDirectory() as temp_dir:
                    # Save input file to temporary location
                    input_path = os.path.join(temp_dir, "input.doc")
                    if isinstance(file_obj, str):
                        with open(file_obj, "rb") as f_src, open(input_path, "wb") as f_dst:
                            f_dst.write(f_src.read())
                    else:
                        file_obj.seek(0)
                        with open(input_path, "wb") as f:
                            f.write(file_obj.read())

                    # Build full path to LibreOffice executable
                    libreoffice_exe = os.path.join("C:\\", "Program Files", "LibreOffice", "program", "soffice.exe")

                    # Verify executable exists
                    if not os.path.exists(libreoffice_exe):
                        raise FileNotFoundError(f"LibreOffice executable not found at: {libreoffice_exe}\nPlease check the installation path.")

                    # Execute conversion command
                    try:
                        cmd = [libreoffice_exe, "--headless", "--convert-to", "docx", "--outdir", temp_dir, input_path]

                        subprocess.run(cmd, capture_output=True, text=True, check=True)

                        # Check for converted file
                        converted_files = [f for f in os.listdir(temp_dir) if f.endswith(".docx")]
                        if not converted_files:
                            raise RuntimeError("LibreOffice conversion produced no output")

                        # Load converted DOCX
                        docx_path = os.path.join(temp_dir, converted_files[0])
                        doc = Document(docx_path)
                        return "\n".join([p.text for p in doc.paragraphs])

                    except subprocess.CalledProcessError as e:
                        logging.error(f"LibreOffice execution failed: {e.stderr}")
                        raise RuntimeError(f"Conversion failed: {e.stderr}") from e

            else:
                raise ValueError(f"Unsupported file type: {ext}")

        except Exception as e:
            logging.error(f"Document processing failed: {str(e)}")
            raise

    def _load_plain_text(self, file_obj: Any) -> str:
        """Plain text loader"""
        try:
            # 注意：file_obj 是二进制文件对象
            content = file_obj.read()

            # 尝试用 UTF-8 解码
            try:
                return content.decode("utf-8")  # type: ignore
            except UnicodeDecodeError:
                # 如果 UTF-8 失败，尝试其他常见编码
                try:
                    return content.decode("latin-1")  # type: ignore
                except Exception as e:
                    raise RuntimeError(f"文本解码失败: {e}") from e
        except Exception as e:
            raise RuntimeError(f"读取文本文件失败: {e}") from e

    def _load_html(self, file_obj: Any) -> str:
        # 读取HTML内容
        html_content = file_obj.read()

        # 常见噪声元素的class/id正则模式
        noise_patterns = [
            r"ad(s|vertisement)?",
            r"banner",
            r"navbar",
            r"footer",
            r"sidebar",
            r"menu",
            r"comment",
            r"widget",
            r"related",
            r"share",
            r"popup",
            r"modal",
            r"cookie",
            r"login",
            r"signup",
            r"subscribe",
            r"promo",
            r"sponsor",
            r"tracking",
        ]

        # 需要直接移除的标签
        remove_tags = ["script", "style", "iframe", "noscript", "svg", "img", "form"]

        soup = BeautifulSoup(html_content, "html.parser")

        # 1. 移除噪声标签
        for tag in remove_tags:
            for element in soup.find_all(tag):
                element.decompose()

        # 2. 移除匹配噪声模式的元素
        for pattern in noise_patterns:
            # 按class移除
            for element in soup.find_all(class_=re.compile(pattern, re.I)):
                element.decompose()
            # 按id移除
            for element in soup.find_all(id=re.compile(pattern, re.I)):
                element.decompose()

        # 3. 获取纯文本并清理
        text = soup.get_text("\n", strip=True)

        # 4. 移除多余空行和首尾空白
        clean_text = re.sub(r"\n{3,}", "\n\n", text).strip()

        return clean_text

    def _load_markdown(self, md_content: str) -> str:
        # Markdown转换规则列表
        transformations = [
            (r"^#+\s*(.*?)\s*$", r"\1"),  # 标题
            (r"\[([^\]]+)\]\([^\)]+\)", r"\1"),  # 链接
            (r"!\[([^\]]*)\]\([^\)]+\)", r"\1"),  # 图片
            (r"`([^`]+)`", r"\1"),  # 内联代码
            (r"```.*?\n(.*?)```", r"\1"),  # 代码块（多行）
            (r"(\*\*|__)(.*?)\1", r"\2"),  # 加粗
            (r"(\*|_)(.*?)\1", r"\2"),  # 斜体
            (r"~~(.*?)~~", r"\1"),  # 删除线
            (r"^[\s]*[-*+]\s+", ""),  # 无序列表
            (r"^[\s]*\d+\.\s+", ""),  # 有序列表
            (r"^>\s*", ""),  # 引用
            (r"\|([^\|]+)\|", r"\1"),  # 表格单元格
            (r"^[-:|]+\s*$", ""),  # 表格分隔线
            (r"<[^>]+>", ""),  # HTML标签
            (r"\n{3,}", "\n\n"),  # 多个空行合并
            (r"[ \t]{2,}", " "),  # 多个空格合并1
        ]

        text = md_content
        for pattern, replacement in transformations:
            text = re.sub(pattern, replacement, text, flags=re.MULTILINE | re.DOTALL)

        return text.strip()

    def _load_json(self, file_obj: Any) -> str:
        """json document loader"""
        text = file_obj.read()
        return text  # type: ignore
